#
#    Copyright (c) 2023 Gérard Parat <aero@lunique.fr>
#
"""Création de bons vol."""

import gnupg
import qrcode
import string
import random
import re
import os
import wx
import platform
import sys

from openpyxl import load_workbook
from docx import Document
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from FloatImageWord import add_float_picture
from docx2pdf import convert

# Case à cocher dans Word
def checkedElement():
    """Coche d'une case à cocher Word."""
    elm = OxmlElement('w:checked')
    elm.set(qn('w:val'),'true')
    return elm

# Génération du numéro de bon
def ID_Generator(maxalpha, maxdigit, chars=string.ascii_uppercase + string.digits):
    """Génération aléatoire de 'maxalpha' lettres et 'maxdigit' chiffres,
les lettres O, I, U et Q ne sont pas autorisées."""
    numeroBon = ''
    nombreTirage = 0
    nombreAlpha = 0
    nombreDigit = 0
    while nombreTirage < maxalpha + maxdigit:
        Tirage = random.choice(chars)
        if Tirage.isdigit() and nombreDigit < maxdigit:
            nombreTirage += 1
            nombreDigit += 1
            numeroBon += Tirage
        elif Tirage.isalpha() and not re.match('(O|I|U|Q)', Tirage) and nombreAlpha < maxalpha:
            nombreTirage += 1
            nombreAlpha += 1
            numeroBon += Tirage
    numeroBonRandom = list(numeroBon)
    random.shuffle(numeroBonRandom)
    return ''.join(numeroBonRandom)

# Calcul de la case en cours
class CaseExcel:
    """Calcul de la case en cours dans un fichier Excel."""

    def __init__(self, colonne,ligne):
        """Case initiale."""
        self.colonne = colonne
        self.ligne = ligne

    def IncrColonne(self):
        """Incrémentation de la colonne jusqu'à ZZ."""
        if len(self.colonne) == 1:
            if self.colonne == 'Z':
                self.colonne = 'AA'
            else:
                self.colonne = chr(ord(self.colonne) + 1)
        elif self.colonne[1] == 'Z':
            self.colonne = chr(ord(self.colonne[0]) + 1) + 'A'
        else:
            self.colonne = self.colonne[0] + chr(ord(self.colonne[1]) + 1)
        return self.colonne + str(self.ligne)

    def IncrLigne(self):
        """Incrémentation de la ligne."""
        self.ligne += 1
        return self.colonne + str(self.ligne)

    def Position(self):
        """Position actuelle de la case."""
        return self.colonne + str(self.ligne)

def genereBon(
        Version, bon, nom1, prenom1, typebon, nom2, prenom2, nom3, prenom3,
        autoparent, choixPaiement, payeur, datePaiement, numcheque, banque,
        date1, heure1, temps1, pilote1, avion1, cours,
        date2, heure2, temps2, pilote2, avion2, tarif,
        cheminVD, classeurVD, modeleVD, cheminVI, classeurVI, modeleVI,
        suiteOffice, cheminLibreOffice, cheminGPG, clef, debug, finTravail):

    """Génération des fichiers de bons de vol."""

    if suiteOffice == 'LibreOffice':
        from unoserver.converter import UnoConverter
    # Spécifique Windows,  pour éviter le conflit
    # avec mailmerge.py de LibreOffice
    if platform.system() == 'Windows':
        if sys.path.count(cheminLibreOffice):
            sys.path.remove(cheminLibreOffice)
        try:
            from mailmerge import MailMerge
        except:
            print('Chemin vers LibreOffice erroné !')
            # Retour
            finTravail[0] = True
            return
    elif platform.system() == 'Linux':
        from mailmerge import MailMerge

    gpg = gnupg.GPG(gpgbinary=cheminGPG)
    gpg.encoding = 'utf-8'

    # Initialisation des fichiers
    if typebon == 'vd':
        ModeleBon = modeleVD
        FichierExcel = classeurVD
        cheminVol = cheminVD
        wb = load_workbook(FichierExcel)
        classeur = wb['vol de découverte']
    else:
        ModeleBon = modeleVI
        FichierExcel = classeurVI
        cheminVol = cheminVI
        wb = load_workbook(FichierExcel)
        classeur = wb['vol d\'initiation']
    
    # Initialisation des données
    NumeroBon = bon
    # Mise en capitale du nom de famille
    Nom1 = str.upper(nom1)
    Nom2 = str.upper(nom2)
    Nom3 = str.upper(nom3)
    # Mise en capitale de la première lettre du prénom
    Prenom1 = str.title(prenom1)
    Prenom2 = str.title(prenom2)
    Prenom3 = str.title(prenom3)
    if autoparent:
        AutoParent = 'OUI'
    else:
        AutoParent = 'NON'
    Paiement = choixPaiement
    NumCheque = numcheque
    Banque = banque
    Payeur = payeur
    if date1 != '' and heure1 != '' and temps1 != '':
        DateVol1 = date1.Format('%d/%m/%Y')
        HeureVol1 = heure1.Format('%Hh%M')
        TempsVol1 = temps1.Format('%Hh%M')
    else:
        DateVol1 = ''
        HeureVol1 = ''
        TempsVol1 = ''
    AvionVol1 = avion1
    PiloteVol1 = pilote1
    if date2 != '' and heure2 != '' and temps2 != '':
        DateVol2 = date2.Format('%d/%m/%Y')
        HeureVol2 = heure2.Format('%Hh%M')
        TempsVol2 = temps2.Format('%Hh%M')
    else:
        DateVol2 = ''
        HeureVol2 = ''
        TempsVol2 = ''
    AvionVol2 = avion2
    PiloteVol2 = pilote2
    
    # Initialisation des dates
    DateFichier = datePaiement.Format('%Y%m%d')
    DateBon = datePaiement.Format('%d/%m/%Y')

    # Composition du texte pour signature
    TexteBonVol = \
    'Numéro de bon : ' + NumeroBon + '\n'
    if typebon == 'vd':
        TexteBonVol += \
        'Passager n° 1 : ' + Nom1 + ' ' + Prenom1 + '\n' + \
        'Passager n° 2 : ' + Nom2 + ' ' + Prenom2 + '\n' + \
        'Passager n° 3 : ' + Nom3 + ' ' + Prenom3 + '\n'
    else:
        TexteBonVol += \
        'Nom : ' + Nom1 + '\n' + \
        'Prénom : ' + Prenom1 + '\n'
    TexteBonVol += \
    'Autorisation parentale : ' + AutoParent + '\n'
    if typebon == 'vi':
        if cours == '1' or cours == '1+1':
            TexteBonVol += \
            'Nombre de vol : 1\n' + \
            'Tarif : ' + str(tarif) + ' EUR\n'
        else:
            TexteBonVol += \
            'Nombre de vol : 2\n' + \
            'Tarif : ' + str(tarif) + ' EUR\n'
    TexteBonVol += \
    'Moyen de Paiement : ' + Paiement + '\n' + \
    'Nom du payeur : ' + Payeur + '\n'
    if Paiement == 'Chèque':
        TexteBonVol += \
        'N° de chèque : ' + NumCheque + '\n' + \
        'Banque : ' + Banque + '\n'
    TexteBonVol += \
    'Date : ' + DateBon + '\n'
    if date1 != '' and heure1 != '' and temps1 != '':
        if typebon == 'vd':
            TexteBonVol += \
            'Date du vol : ' + DateVol1 + '\n'
        elif cours == '1+1':
            TexteBonVol += \
            'Date du vol n° 2 : ' + DateVol1 + '\n'
        else:
            TexteBonVol += \
            'Date du vol n° 1 : ' + DateVol1 + '\n'
        TexteBonVol += \
        'Heure du vol : ' + HeureVol1 + '\n' + \
        'Durée du vol : ' + TempsVol1 + '\n'
    if AvionVol1 != '' and PiloteVol1 != '':
        TexteBonVol += \
        'Avion : ' + AvionVol1 + '\n' + \
        'Pilote : ' + PiloteVol1 + '\n'
    if date2 != '' and heure2 != '' and temps2 != '':
        TexteBonVol += \
        'Date du vol n° 2 : ' + DateVol2 + '\n' + \
        'Heure du vol : ' + HeureVol2 + '\n' + \
        'Durée du vol : ' + TempsVol2 + '\n'
    if AvionVol2 != '' and PiloteVol2 != '':
        TexteBonVol += \
        'Avion : ' + AvionVol2 + '\n' + \
        'Pilote : ' + PiloteVol2 + '\n'

    # Nom des fichiers
    NomFichier = os.path.join(cheminVol, DateFichier + '-' + Nom1 + '-' + Prenom1 + '-' + NumeroBon)

    # Signature authentique du texte avec une clef
    signature_filename = NomFichier + '.asc'
    gpg.sign(TexteBonVol, keyid=clef, output=signature_filename)

    # Erreur si la phrase secrète de la clef est erronée ou manquante
    gpg.encoding = 'latin-1'
    try:
        with open(signature_filename, 'rb') as signature_file:
            verified = gpg.verify_file(signature_file)
    except FileNotFoundError:
        print('Le fichier %s n\'a pas été créé !' % signature_filename,
              'Phrase secrète erronée ou manquante...',
              '**TEMPORISATION** de 30 s', sep=os.linesep)
        # Retour
        finTravail[0] = True
        return
    else:
        if not verified:
            if not debug:
                os.remove(signature_filename)
            print('Problème de signature !')
            # Retour
            finTravail[0] = True
            return
    
    # Génération du QR-Code
    with open(signature_filename, 'rb') as signature_file:
        QRimage = qrcode.make(signature_file.read())
        QRimage.save(NomFichier + '.png')
    
    # Remplissage du fichier Excel des bons de vol
    # La première case à remplir est en A7
    case = CaseExcel('A',7)
    
    # Recherche de la première ligne vide
    while classeur[case.Position()].value != None:
        case.IncrLigne()
    else:
        # Remplissage des valeurs sur la première ligne libre
        if typebon == 'vd':
            classeur[case.Position()] = NumeroBon
            classeur[case.IncrColonne()] = Nom1 + ' ' + Prenom1
            classeur[case.IncrColonne()] = Nom2 + ' ' + Prenom2
            classeur[case.IncrColonne()] = Nom3 + ' ' + Prenom3
            classeur[case.IncrColonne()] = Paiement
            classeur[case.IncrColonne()] = Payeur
            classeur[case.IncrColonne()] = DateBon
            classeur[case.IncrColonne()] = DateVol1
            classeur[case.IncrColonne()] = HeureVol1
            classeur[case.IncrColonne()] = TempsVol1
            classeur[case.IncrColonne()] = AvionVol1
            classeur[case.IncrColonne()] = PiloteVol1
            classeur[case.IncrColonne()] = AutoParent
        else:
            classeur[case.Position()] = NumeroBon
            classeur[case.IncrColonne()] = Nom1 + ' ' + Prenom1
            if cours == '1' or cours == '1+1':
                classeur[case.IncrColonne()] = 1
            else:
                classeur[case.IncrColonne()] = 2
            classeur[case.IncrColonne()] = tarif
            classeur[case.IncrColonne()] = Paiement
            classeur[case.IncrColonne()] = Payeur
            classeur[case.IncrColonne()] = DateBon
            classeur[case.IncrColonne()] = DateVol1
            classeur[case.IncrColonne()] = AvionVol1
            classeur[case.IncrColonne()] = PiloteVol1
            if cours == '2':
                classeur[case.IncrColonne()] = DateVol2
                classeur[case.IncrColonne()] = AvionVol2
                classeur[case.IncrColonne()] = PiloteVol2
            else:
                case.IncrColonne()
                case.IncrColonne()
                case.IncrColonne()
            case.IncrColonne()
            classeur[case.IncrColonne()] = AutoParent

    # Ecriture du fichier Excel des bons
    wb.save(filename = FichierExcel)
    
    # Création du bon
    DateValide = datePaiement.Add(wx.DateSpan(1))
    with MailMerge(ModeleBon) as BonVol:
        if typebon == 'vd':
            BonVol.merge(
                NuméroBon = NumeroBon,
                Pax1 = Nom1 + ' ' + Prenom1,
                Pax2 = Nom2 + ' ' + Prenom2,
                Pax3 = Nom3 + ' ' + Prenom3,
                DateRèglement = DateBon,
                NomPayeur = Payeur,
                NuméroChèque = NumCheque,
                NomBanque = Banque,
                DateValidité = DateValide.Format('%d/%m/%Y'),
                Pilote = PiloteVol1,
                Avion = AvionVol1,
                DateVol = DateVol1,
                HeureVol = HeureVol1,
                DuréeVol = TempsVol1,
                Version = Version)
        else:
            BonVol.merge(
                NuméroBon = NumeroBon,
                NomPrénom = Nom1 + ' ' + Prenom1,
                Instr1 = PiloteVol1,
                Avion1 = AvionVol1,
                Date1 = DateVol1,
                Heure1 = HeureVol1,
                Durée1 = TempsVol1,
                Instr2 = PiloteVol2,
                Avion2 = AvionVol2,
                Date2 = DateVol2,
                Heure2 = HeureVol2,
                Durée2 = TempsVol2,
                DateRèglement = DateBon,
                NomPayeur = Payeur,
                NuméroChèque = NumCheque,
                NomBanque = Banque,
                DateValidité = DateValide.Format('%d/%m/%Y'),
                Version = Version)
        BonVol.write(NomFichier + '.docx')

    # Ajout du QR-Code
    BonVol = Document(NomFichier + '.docx')
    tables = BonVol.tables
    p = BonVol.tables[0].rows[0].cells[1].add_paragraph()
    add_float_picture(p, NomFichier + '.png', width=Cm(6.5))
    # Cases à cocher
    checkboxes = BonVol._element.xpath('.//w:checkBox')
    if typebon == 'vd':
        if AutoParent == 'OUI':
            checkboxes[0].append(checkedElement())
        if Paiement == 'Espèce':
            checkboxes[1].append(checkedElement())
        elif Paiement == 'Carte bancaire':
            checkboxes[2].append(checkedElement())
        elif Paiement == 'Virement':
            checkboxes[3].append(checkedElement())
        else:
            checkboxes[4].append(checkedElement())
    else:
        if AutoParent == 'OUI':
            checkboxes[0].append(checkedElement())
        if Paiement == 'Espèce':
            checkboxes[1].append(checkedElement())
        elif Paiement == 'Carte bancaire':
            checkboxes[2].append(checkedElement())
        elif Paiement == 'Virement':
            checkboxes[3].append(checkedElement())
        else:
            checkboxes[4].append(checkedElement())
        if cours == '2':
            checkboxes[6].append(checkedElement())
    BonVol.save(NomFichier + '.docx')

    # Création du PDF avec Microsoft Office (Windows)
    if platform.system() == 'Windows' and suiteOffice == 'Microsoft Office':
        convert(NomFichier + '.docx')
    # Création du PDF avec LibreOffice (Windows ou Linux)
    elif suiteOffice == 'LibreOffice':
        convertisseur = UnoConverter()
        convertisseur.convert(inpath=NomFichier + '.docx', outpath=NomFichier + '.pdf')

    # Nettoyage éventuel des fichiers intermédiaires
    if not debug:
        os.remove(NomFichier + '.asc')
        os.remove(NomFichier + '.png')
        os.remove(NomFichier + '.docx')
    
    # Retour
    finTravail[0] = True
    finTravail[1] = True
